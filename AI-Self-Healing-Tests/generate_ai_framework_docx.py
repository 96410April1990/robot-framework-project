#!/usr/bin/env python3
"""
AI Self-Healing Framework Documentation Generator
Converts Markdown to DOCX format with proper formatting
"""
import os
import sys
from pathlib import Path

try:
    import markdown
    from markdown.extensions import tables
    import pdfkit
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    import re
except ImportError as e:
    print(f"Missing required library: {e}")
    print("Installing required packages...")
    os.system("pip install markdown pdfkit beautifulsoup4 python-docx")
    import markdown
    from markdown.extensions import tables
    import pdfkit
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    import re

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set hyperlink style
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)

def setup_document_styles(doc):
    """Setup custom styles for the document"""
    styles = doc.styles
    
    # Title style
    if 'Custom Title' not in [s.name for s in styles]:
        title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(24)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    if 'Custom Heading 1' not in [s.name for s in styles]:
        h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Calibri'
        h1_font.size = Pt(18)
        h1_font.bold = True
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    if 'Custom Heading 2' not in [s.name for s in styles]:
        h2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Calibri'
        h2_font.size = Pt(16)
        h2_font.bold = True
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(4)
    
    # Heading 3 style
    if 'Custom Heading 3' not in [s.name for s in styles]:
        h3_style = styles.add_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = 'Calibri'
        h3_font.size = Pt(14)
        h3_font.bold = True
        h3_style.paragraph_format.space_before = Pt(8)
        h3_style.paragraph_format.space_after = Pt(4)
    
    # Code style
    if 'Custom Code' not in [s.name for s in styles]:
        code_style = styles.add_style('Custom Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Consolas'
        code_font.size = Pt(10)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)

def process_markdown_line(doc, line, in_code_block=False):
    """Process a single line of markdown and add to document"""
    line = line.rstrip()
    
    # Skip empty lines in code blocks
    if in_code_block and not line.strip():
        return in_code_block
    
    # Code block detection
    if line.startswith('```'):
        return not in_code_block
    
    if in_code_block:
        # Add code line
        para = doc.add_paragraph(line, style='Custom Code')
        return in_code_block
    
    # Headers
    if line.startswith('# '):
        para = doc.add_paragraph(line[2:], style='Custom Title')
    elif line.startswith('## '):
        para = doc.add_paragraph(line[3:], style='Custom Heading 1')
    elif line.startswith('### '):
        para = doc.add_paragraph(line[4:], style='Custom Heading 2')
    elif line.startswith('#### '):
        para = doc.add_paragraph(line[5:], style='Custom Heading 3')
    
    # Table detection (simple)
    elif '|' in line and line.count('|') >= 2:
        # Skip table separator lines
        if re.match(r'^\|[\s\-\|\:]+\|$', line):
            return in_code_block
        
        # Process table row
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        
        # Create table if this is first row
        if not hasattr(process_markdown_line, 'current_table'):
            process_markdown_line.current_table = doc.add_table(rows=1, cols=len(cells))
            process_markdown_line.current_table.style = 'Table Grid'
            
            # Add header row
            hdr_cells = process_markdown_line.current_table.rows[0].cells
            for i, cell_text in enumerate(cells):
                hdr_cells[i].text = cell_text.strip('*')
                # Make header bold
                for para in hdr_cells[i].paragraphs:
                    for run in para.runs:
                        run.bold = True
        else:
            # Add data row
            row_cells = process_markdown_line.current_table.add_row().cells
            for i, cell_text in enumerate(cells):
                if i < len(row_cells):
                    row_cells[i].text = cell_text.strip('*')
    
    # Lists
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        para = doc.add_paragraph(text, style='List Bullet')
    elif re.match(r'^\d+\. ', line):
        text = re.sub(r'^\d+\. ', '', line).strip()
        para = doc.add_paragraph(text, style='List Number')
    
    # Indented lists
    elif line.startswith('├─ ') or line.startswith('└─ '):
        text = line[3:].strip()
        para = doc.add_paragraph(text, style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.5)
    
    # Bold/emphasis handling for regular paragraphs
    elif line.strip():
        # Reset table tracking for new content
        if hasattr(process_markdown_line, 'current_table'):
            delattr(process_markdown_line, 'current_table')
        
        para = doc.add_paragraph()
        
        # Process text with bold markers
        text = line.strip()
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part:
                para.add_run(part)
    
    return in_code_block

def markdown_to_docx(markdown_file, output_file):
    """Convert markdown file to DOCX format"""
    
    # Read markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Setup styles
    setup_document_styles(doc)
    
    # Process content line by line
    lines = content.split('\n')
    in_code_block = False
    
    for line in lines:
        in_code_block = process_markdown_line(doc, line, in_code_block)
    
    # Save document
    doc.save(output_file)
    print(f"Successfully converted to DOCX: {output_file}")

def main():
    """Main function to handle the conversion"""
    
    # File paths
    current_dir = Path(__file__).parent
    ai_folder = current_dir
    
    markdown_file = ai_folder / "AI_Self_Healing_Framework_Overview.md"
    docx_file = ai_folder / "AI_Self_Healing_Framework_Overview.docx"
    
    # Check if markdown file exists
    if not markdown_file.exists():
        print(f"Error: Markdown file not found at {markdown_file}")
        return 1
    
    # Convert to DOCX
    try:
        markdown_to_docx(markdown_file, docx_file)
        
        print("\n" + "="*60)
        print("AI SELF-HEALING FRAMEWORK DOCUMENTATION GENERATED")
        print("="*60)
        print(f"📁 Location: {ai_folder}")
        print(f"📄 DOCX File: {docx_file.name}")
        print(f"📊 Size: {docx_file.stat().st_size / 1024:.1f} KB")
        print("\n✅ Document ready for stakeholder distribution!")
        print("="*60)
        
        return 0
        
    except Exception as e:
        print(f"Error during conversion: {e}")
        return 1

if __name__ == "__main__":
    sys.exit(main())